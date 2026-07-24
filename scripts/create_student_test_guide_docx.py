from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_PATH = OUT_DIR / "GAKUZAI_student_test_guide_ja.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="CBD5E1", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_column_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_run(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_note_box(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 9360)
    set_table_borders(table, color="BFDBFE", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EFF6FF")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, bold=True, color="1E3A8A")
    for line in lines:
        lp = cell.add_paragraph(style=None)
        lp.paragraph_format.space_after = Pt(2)
        lp.add_run(line)
    doc.add_paragraph()


def add_info_table(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 9360)
    set_table_borders(table)
    set_column_widths(table, [1.55, 4.95])
    rows = [
        ("接続先URL", "http://　　　　　　　　　　　　:3000/"),
        ("授業コード", "DIGI2026"),
        ("使用端末", "PC、タブレット、スマートフォン"),
        ("注意", "学校のWi-Fiに接続してからアクセスしてください。"),
    ]
    for i, (label, value) in enumerate(rows):
        left, right = table.rows[i].cells
        set_cell_shading(left, "E8EEF5")
        left.paragraphs[0].add_run(label).bold = True
        right.paragraphs[0].add_run(value)
        if label == "接続先URL":
            for run in right.paragraphs[0].runs:
                run.bold = True
        if label == "授業コード":
            for run in right.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(13)
    doc.add_paragraph()


def add_button_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 9360)
    set_table_borders(table)
    set_column_widths(table, [2.0, 4.5])
    hdr = table.rows[0].cells
    hdr[0].text = "ボタン・機能"
    hdr[1].text = "使い方"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    rows = [
        ("マーカー：黄／緑／ピンク", "重要だと思う語句や文に色を付けます。"),
        ("文字色：青／赤", "意味のまとまりや注意点を色で区別します。"),
        ("強調：太字／下線", "特に大切だと思う語句を目立たせます。"),
        ("非表示・キーワード化", "語句を隠したり、短い記号やキーワードに置き換えたりします。"),
        ("ポップアップ追加", "補足説明や自分のメモを語句に付けます。"),
        ("装飾解除", "選択した部分の色、太字、下線、ポップアップなどを元に戻します。"),
        ("保存", "自分の加工結果を保存します。終了前に必ず押してください。"),
    ]
    for label, desc in rows:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].paragraphs[0].add_run(desc)
        for cell in cells:
            set_cell_margins(cell)
    doc.add_paragraph()


def add_shortcut_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 9360)
    set_table_borders(table)
    set_column_widths(table, [2.15, 4.35])
    table.rows[0].cells[0].text = "操作"
    table.rows[0].cells[1].text = "意味"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    rows = [
        ("Esc", "選択中の加工モードを解除します。"),
        ("Ctrl + Z", "一つ前の状態に戻します。Macの場合は Cmd + Z です。"),
        ("Ctrl + Shift + Z", "取り消した操作をやり直します。Macの場合は Cmd + Shift + Z です。"),
        ("Ctrl + Y", "やり直しとして使える場合があります。"),
    ]
    for label, desc in rows:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].paragraphs[0].add_run(desc)
        for cell in cells:
            set_cell_margins(cell)
    doc.add_paragraph()


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color, before, after in [
        ("Heading 1", 15, "2E74B5", 12, 6),
        ("Heading 2", 12.5, "2E74B5", 9, 4),
        ("Heading 3", 11.5, "1F4D78", 6, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Yu Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "GAKUZAI 授業テスト用 簡易操作説明", bold=True, color="0B2545", size=18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    add_run(subtitle, "学生配布用：ログイン、授業参加、教材加工、保存の手順", color="555555", size=10.5)

    add_info_table(doc)

    doc.add_heading("1. はじめに", level=1)
    p = doc.add_paragraph()
    p.add_run("この説明書は、授業中に GAKUZAI を使って電子教材を加工するための簡単な手順です。")
    p.add_run("教材を読んで、自分が重要だと思う部分にマーカー、色、注釈などを付けてください。")

    doc.add_heading("2. ログインまたは新規登録", level=1)
    add_number(doc, "学校の Wi-Fi に接続します。")
    add_number(doc, "上の「接続先URL」をブラウザに入力して開きます。")
    add_number(doc, "初めて使う人は、名前、メールアドレス、パスワードを入力して「登録」を押します。")
    add_number(doc, "すでに登録済みの人は、メールアドレスとパスワードを入力して「ログイン」を押します。")
    add_note_box(doc, "入力時の注意", [
        "メールアドレスとパスワードは、あとで同じ学習記録を開くために使います。",
        "他の人のアカウントではログインしないでください。"
    ])

    doc.add_heading("3. 授業に参加して教材を開く", level=1)
    add_number(doc, "ログイン後、左側またはメニューの「授業・教材管理」を開きます。")
    add_number(doc, "「授業コード」に DIGI2026 と入力します。")
    add_number(doc, "「授業に参加」を押します。")
    add_number(doc, "参加した授業の中から教材を選び、「開く」または「教材加工」を押します。")
    add_number(doc, "教材が表示されたら、内容を読みながら加工を始めます。")

    doc.add_heading("4. 教材加工の基本操作", level=1)
    p = doc.add_paragraph()
    p.add_run("基本操作は、")
    p.add_run("ツールを選ぶ → 加工したい語句をドラッグして選択する").bold = True
    p.add_run("、という順番です。")
    add_button_table(doc)

    doc.add_heading("5. ショートカットキー", level=1)
    add_shortcut_table(doc)

    doc.add_heading("6. スマートフォン・タブレットで使う場合", level=1)
    add_bullet(doc, "画面下部またはメニューから加工ツールを開きます。")
    add_bullet(doc, "語句を長押しして選択し、表示された操作ボタンで適用します。")
    add_bullet(doc, "左下などに表示される「↶」ボタンは、元に戻す操作です。")
    add_bullet(doc, "画面が狭い場合は、横向きにすると操作しやすくなります。")

    doc.add_heading("7. 保存と終了", level=1)
    add_number(doc, "加工が終わったら、必ず「保存」を押します。")
    add_number(doc, "保存後、「学習記録」から自分の加工結果を確認できます。")
    add_number(doc, "先生から課題の指示がある場合は、「課題」画面を開いて回答します。")
    add_number(doc, "最後に「ログアウト」を押して終了します。")
    add_note_box(doc, "授業中に困ったとき", [
        "画面が止まったように見える場合は、すぐに何度も押さず、数秒待ってください。",
        "保存できたか不安な場合は、先生または補助者に声をかけてください。"
    ])

    doc.add_heading("8. 今日のテストでお願いしたいこと", level=1)
    add_bullet(doc, "正解を作ることよりも、自分が理解しやすくなるように教材を加工してください。")
    add_bullet(doc, "重要語句、わかりにくい語句、あとで復習したい部分が分かるように印を付けてください。")
    add_bullet(doc, "他の人と同じ加工にする必要はありません。自分の考えで操作してください。")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("GAKUZAI 授業テスト用説明書").font.size = Pt(9)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_doc()
